# -*- coding: utf-8 -*-
# Template Generator - utils/template_generator.py
# CV és motivációs levél sablonok + .docx generálás

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import os
from datetime import datetime

TEMPLATES_DIR = "templates"


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── CV SABLONOK ───────────────────────────────────────────────

def create_modern_template():
    """Modern, tech pozíciókhoz - kék accent szín"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    nev = doc.add_paragraph()
    nev.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = nev.add_run("{{NEV}}")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    kontakt = doc.add_paragraph("{{VAROS}}  |  {{TELEFON}}  |  {{EMAIL}}")
    kontakt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kontakt.runs[0].font.size = Pt(10)
    kontakt.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    add_horizontal_line(kontakt)

    for cim, tartalom in [
        ("PROFIL", "{{PROFIL}}"),
        ("TAPASZTALAT", "{{TAPASZTALAT}}"),
        ("TANULMÁNYOK", "{{TANULMANYOK}}"),
        ("KÉSZSÉGEK", "{{KESZSEGEK}}"),
        ("PROJEKTEK", "{{PROJEKTEK}}"),
    ]:
        h = doc.add_paragraph()
        h_run = h.add_run(cim)
        h_run.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        add_horizontal_line(h)
        t = doc.add_paragraph(tartalom)
        t.runs[0].font.size = Pt(10)

    doc.save(os.path.join(TEMPLATES_DIR, "modern.docx"))
    print("modern.docx kész")


def create_klasszikus_template():
    """Klasszikus, pénzügyi/jogi pozíciókhoz - fekete/szürke"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    nev = doc.add_paragraph()
    nev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = nev.add_run("{{NEV}}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    kontakt = doc.add_paragraph("{{VAROS}}  •  {{TELEFON}}  •  {{EMAIL}}")
    kontakt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kontakt.runs[0].font.size = Pt(10)
    add_horizontal_line(kontakt)

    for cim, tartalom in [
        ("SZAKMAI ÖSSZEFOGLALÓ", "{{PROFIL}}"),
        ("SZAKMAI TAPASZTALAT", "{{TAPASZTALAT}}"),
        ("VÉGZETTSÉG", "{{TANULMANYOK}}"),
        ("KOMPETENCIÁK", "{{KESZSEGEK}}"),
    ]:
        h = doc.add_paragraph()
        h_run = h.add_run(cim)
        h_run.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        add_horizontal_line(h)
        t = doc.add_paragraph(tartalom)
        t.runs[0].font.size = Pt(10)

    doc.save(os.path.join(TEMPLATES_DIR, "klasszikus.docx"))
    print("klasszikus.docx kész")


def create_kreativ_template():
    """Kreatív, marketing/design pozíciókhoz - lila accent"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    nev = doc.add_paragraph()
    nev.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = nev.add_run("{{NEV}}")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x7B, 0x2D, 0xBF)

    pozicio = doc.add_paragraph("{{POZICIO_CIM}}")
    pozicio.runs[0].font.size = Pt(13)
    pozicio.runs[0].font.color.rgb = RGBColor(0x7B, 0x2D, 0xBF)
    pozicio.runs[0].italic = True

    kontakt = doc.add_paragraph("{{VAROS}}  |  {{TELEFON}}  |  {{EMAIL}}")
    kontakt.runs[0].font.size = Pt(10)
    add_horizontal_line(kontakt)

    for cim, tartalom in [
        ("RÓLAM", "{{PROFIL}}"),
        ("TAPASZTALAT", "{{TAPASZTALAT}}"),
        ("TANULMÁNYOK", "{{TANULMANYOK}}"),
        ("KÉSZSÉGEK & ESZKÖZÖK", "{{KESZSEGEK}}"),
        ("PROJEKTEK & PORTFÓLIÓ", "{{PROJEKTEK}}"),
    ]:
        h = doc.add_paragraph()
        h_run = h.add_run(cim)
        h_run.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(0x7B, 0x2D, 0xBF)
        add_horizontal_line(h)
        t = doc.add_paragraph(tartalom)
        t.runs[0].font.size = Pt(10)

    doc.save(os.path.join(TEMPLATES_DIR, "kreativ.docx"))
    print("kreativ.docx kész")


# ── MOTIVÁCIÓS LEVÉL SABLONOK ─────────────────────────────────

def create_motivacios_modern():
    """Modern motivációs levél — tech/IT pozíciókhoz"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Fejléc — név és kontakt
    nev = doc.add_paragraph()
    run = nev.add_run("{{NEV}}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    kontakt = doc.add_paragraph("{{VAROS}}  |  {{TELEFON}}  |  {{EMAIL}}")
    kontakt.runs[0].font.size = Pt(10)
    kontakt.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    add_horizontal_line(kontakt)

    # Dátum
    datum = doc.add_paragraph()
    datum.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d_run = datum.add_run(datetime.now().strftime("%Y. %B %d."))
    d_run.font.size = Pt(10)
    d_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Tárgy
    targy = doc.add_paragraph()
    t_run = targy.add_run("Tárgy: Jelentkezés — {{POZICIO}} pozícióra")
    t_run.bold = True
    t_run.font.size = Pt(11)
    t_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    doc.add_paragraph()

    # Levél törzse
    level = doc.add_paragraph("{{LEVEL_SZOVEGE}}")
    level.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Aláírás
    alairas = doc.add_paragraph()
    a_run = alairas.add_run("Üdvözlettel,\n{{NEV}}")
    a_run.font.size = Pt(11)

    doc.save(os.path.join(TEMPLATES_DIR, "motivacios_modern.docx"))
    print("motivacios_modern.docx kész")


def create_motivacios_klasszikus():
    """Klasszikus motivációs levél — pénzügyi/jogi pozíciókhoz"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Fejléc
    nev = doc.add_paragraph()
    nev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = nev.add_run("{{NEV}}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    kontakt = doc.add_paragraph("{{VAROS}}  •  {{TELEFON}}  •  {{EMAIL}}")
    kontakt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kontakt.runs[0].font.size = Pt(10)
    add_horizontal_line(kontakt)

    doc.add_paragraph()

    # Dátum jobbra
    datum = doc.add_paragraph()
    datum.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d_run = datum.add_run(datetime.now().strftime("%Y. %B %d."))
    d_run.font.size = Pt(10)

    doc.add_paragraph()

    # Megszólítás
    megszolitas = doc.add_paragraph("Tisztelt {{CEG}}!")
    megszolitas.runs[0].font.size = Pt(11)
    megszolitas.runs[0].bold = True

    doc.add_paragraph()

    # Levél törzse
    level = doc.add_paragraph("{{LEVEL_SZOVEGE}}")
    level.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Aláírás
    alairas = doc.add_paragraph("Tisztelettel,")
    alairas.runs[0].font.size = Pt(11)

    nev_alairas = doc.add_paragraph()
    n_run = nev_alairas.add_run("{{NEV}}")
    n_run.bold = True
    n_run.font.size = Pt(11)

    doc.save(os.path.join(TEMPLATES_DIR, "motivacios_klasszikus.docx"))
    print("motivacios_klasszikus.docx kész")


def create_motivacios_kreativ():
    """Kreatív motivációs levél — marketing/design pozíciókhoz"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Fejléc
    nev = doc.add_paragraph()
    run = nev.add_run("{{NEV}}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x7B, 0x2D, 0xBF)

    pozicio = doc.add_paragraph("{{POZICIO}} pozícióra jelentkezem")
    pozicio.runs[0].font.size = Pt(12)
    pozicio.runs[0].italic = True
    pozicio.runs[0].font.color.rgb = RGBColor(0x7B, 0x2D, 0xBF)

    kontakt = doc.add_paragraph("{{VAROS}}  |  {{TELEFON}}  |  {{EMAIL}}")
    kontakt.runs[0].font.size = Pt(10)
    add_horizontal_line(kontakt)

    doc.add_paragraph()

    # Levél törzse
    level = doc.add_paragraph("{{LEVEL_SZOVEGE}}")
    level.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Aláírás
    alairas = doc.add_paragraph()
    a_run = alairas.add_run("{{NEV}}")
    a_run.bold = True
    a_run.font.size = Pt(11)
    a_run.font.color.rgb = RGBColor(0x7B, 0x2D, 0xBF)

    doc.save(os.path.join(TEMPLATES_DIR, "motivacios_kreativ.docx"))
    print("motivacios_kreativ.docx kész")


# ── .DOCX GENERÁLÁS — KITÖLTÖTT FÁJLOK ───────────────────────

def generalt_cv_docx(cv_szoveg: str, sablon: str = "Modern",
                      nev: str = "", varos: str = "",
                      telefon: str = "", email: str = "") -> BytesIO:
    """Kitölt egy CV sablont és visszaadja BytesIO-ként letöltéshez"""

    sablon_fajl = {
        "Modern": "modern.docx",
        "Klasszikus": "klasszikus.docx",
        "Kreatív": "kreativ.docx"
    }.get(sablon, "modern.docx")

    sablon_ut = os.path.join(TEMPLATES_DIR, sablon_fajl)

    if not os.path.exists(sablon_ut):
        run()  # Generáljuk a sablonokat ha nem léteznek

    doc = Document(sablon_ut)

    # Placeholderek cseréje
    cserék = {
        "{{NEV}}": nev or "Varga Andrea",
        "{{VAROS}}": varos or "Székesfehérvár",
        "{{TELEFON}}": telefon or "+36 70 358 7697",
        "{{EMAIL}}": email or "varga.andrea.job@gmail.com",
        "{{PROFIL}}": cv_szoveg,
        "{{TAPASZTALAT}}": "",
        "{{TANULMANYOK}}": "",
        "{{KESZSEGEK}}": "",
        "{{PROJEKTEK}}": "",
        "{{POZICIO_CIM}}": "",
    }

    for para in doc.paragraphs:
        for kulcs, ertek in cserék.items():
            if kulcs in para.text:
                for run in para.runs:
                    if kulcs in run.text:
                        run.text = run.text.replace(kulcs, ertek)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generalt_motivacios_docx(level_szoveg: str, sablon: str = "Modern",
                              nev: str = "", varos: str = "",
                              telefon: str = "", email: str = "",
                              pozicio: str = "", ceg: str = "") -> BytesIO:
    """Kitölt egy motivációs levél sablont és visszaadja BytesIO-ként"""

    sablon_fajl = {
        "Modern": "motivacios_modern.docx",
        "Klasszikus": "motivacios_klasszikus.docx",
        "Kreatív": "motivacios_kreativ.docx"
    }.get(sablon, "motivacios_modern.docx")

    sablon_ut = os.path.join(TEMPLATES_DIR, sablon_fajl)

    if not os.path.exists(sablon_ut):
        run()

    doc = Document(sablon_ut)

    cserék = {
        "{{NEV}}": nev or "Varga Andrea",
        "{{VAROS}}": varos or "Székesfehérvár",
        "{{TELEFON}}": telefon or "+36 70 358 7697",
        "{{EMAIL}}": email or "varga.andrea.job@gmail.com",
        "{{POZICIO}}": pozicio,
        "{{CEG}}": ceg or "Tisztelt Cég",
        "{{LEVEL_SZOVEGE}}": level_szoveg,
    }

    for para in doc.paragraphs:
        for kulcs, ertek in cserék.items():
            if kulcs in para.text:
                for run in para.runs:
                    if kulcs in run.text:
                        run.text = run.text.replace(kulcs, ertek)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ── SABLONOK GENERÁLÁSA ───────────────────────────────────────

def run():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    create_modern_template()
    create_klasszikus_template()
    create_kreativ_template()
    create_motivacios_modern()
    create_motivacios_klasszikus()
    create_motivacios_kreativ()
    print("Összes sablon elkészült!")


if __name__ == "__main__":
    run()