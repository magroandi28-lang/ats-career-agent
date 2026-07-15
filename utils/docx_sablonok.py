# -*- coding: utf-8 -*-
# Word (.docx) Sablon Generátor - utils/docx_sablonok.py
# Elegáns, visszafogott, ATS-barát CV + motivációs levél (python-docx)

import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SZINEK = {
    "kek":     {"fo": "1a3a5c", "akcent": "2563eb", "nev": "Kék"},
    "arany":   {"fo": "1a1a2e", "akcent": "b8860b", "nev": "Arany"},
    "zold":    {"fo": "14321e", "akcent": "16a34a", "nev": "Zöld"},
    "bordó":   {"fo": "3a1620", "akcent": "be123c", "nev": "Bordó"},
    "antracit":{"fo": "1f2937", "akcent": "475569", "nev": "Antracit"},
}

FEKETE   = RGBColor(0x1f, 0x29, 0x37)
SOTET    = RGBColor(0x11, 0x18, 0x27)
SZURKE   = RGBColor(0x6b, 0x72, 0x80)
HAIRLINE = "d8dee9"


def szin_ajanlat(szakma_kategoria: str) -> str:
    if not szakma_kategoria:
        return "arany"
    kat = szakma_kategoria.lower()
    if any(s in kat for s in ["it", "informatika", "programozó", "fejlesztő", "mérnök"]):
        return "kek"
    if any(s in kat for s in ["egészség", "ápoló", "orvos"]):
        return "zold"
    if any(s in kat for s in ["jog", "pénzügy", "könyvelő"]):
        return "bordó"
    if any(s in kat for s in ["marketing", "kreatív"]):
        return "antracit"
    return "arany"


def szinek_listaja() -> list:
    return [
        {"kulcs": "kek",      "nev": "🔵 Kék",      "leiras": "IT, tech"},
        {"kulcs": "arany",    "nev": "🟡 Arany",    "leiras": "Üzleti, elegáns"},
        {"kulcs": "zold",     "nev": "🟢 Zöld",     "leiras": "Egészségügy, szakma"},
        {"kulcs": "bordó",    "nev": "🔴 Bordó",    "leiras": "Jog, pénzügy"},
        {"kulcs": "antracit", "nev": "⚫ Antracit", "leiras": "Modern, semleges"},
    ]


def _szegely(p, oldal, color_hex, meret):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        pPr.append(pbdr)
    el = OxmlElement(f'w:{oldal}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(meret))
    el.set(qn('w:space'), '6')
    el.set(qn('w:color'), color_hex)
    pbdr.append(el)


def _betuk_ritkitas(run, ertek=20):
    rPr = run._element.get_or_add_rPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:val'), str(ertek))
    rPr.append(sp)


def _alap_dokumentum():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = FEKETE
    pf = style.paragraph_format
    pf.line_spacing = 1.22
    pf.space_after = Pt(0)

    sz = doc.sections[0]
    sz.page_height = Cm(29.7)
    sz.page_width = Cm(21.0)
    sz.top_margin = Cm(1.6)
    sz.bottom_margin = Cm(1.6)
    sz.left_margin = Cm(2.0)
    sz.right_margin = Cm(2.0)
    return doc
def _fejlec(doc, adatok, fo_hex, akcent_hex):
    nev = adatok.get("nev", "")
    pozicio = adatok.get("pozicio", "")
    email = adatok.get("email", "")
    telefon = adatok.get("telefon", "")
    varos = adatok.get("varos", "")

    p_nev = doc.add_paragraph()
    p_nev.paragraph_format.space_before = Pt(0)
    p_nev.paragraph_format.space_after = Pt(2)
    _szegely(p_nev, 'top', akcent_hex, 24)
    r = p_nev.add_run(nev)
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = SOTET
    r.font.name = 'Arial'

    if pozicio:
        p_poz = doc.add_paragraph()
        p_poz.paragraph_format.space_after = Pt(4)
        rp = p_poz.add_run(pozicio.upper())
        rp.font.size = Pt(10.5)
        rp.font.color.rgb = RGBColor.from_string(akcent_hex)
        rp.font.name = 'Arial'
        _betuk_ritkitas(rp, 40)

    reszek = [x for x in (telefon, email, varos) if x]
    if reszek:
        p_el = doc.add_paragraph()
        p_el.paragraph_format.space_after = Pt(4)
        re = p_el.add_run("   ·   ".join(reszek))
        re.font.size = Pt(9)
        re.font.color.rgb = SZURKE
        re.font.name = 'Arial'
        _szegely(p_el, 'bottom', HAIRLINE, 6)


def _profil_blokk(doc, sorok, akcent_hex):
    szoveg = " ".join(s.strip() for s in sorok if s.strip())
    if not szoveg:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    _szegely(p, 'left', akcent_hex, 18)
    r = p.add_run(szoveg.replace('**', ''))
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


def _szekcio_cim(doc, szoveg, akcent_hex):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    marker = p.add_run("▪  ")
    marker.font.size = Pt(10)
    marker.font.color.rgb = RGBColor.from_string(akcent_hex)
    marker.font.name = 'Arial'
    r = p.add_run(szoveg.upper())
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = 'Arial'
    r.font.color.rgb = SOTET
    _betuk_ritkitas(r, 30)
    _szegely(p, 'bottom', HAIRLINE, 6)


def _munka_cim(doc, szoveg):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(szoveg)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = 'Arial'
    r.font.color.rgb = FEKETE


def _munka_ceg(doc, szoveg):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(szoveg)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = 'Arial'
    r.font.color.rgb = SZURKE


def _bekezdes(doc, szoveg):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(szoveg)
    r.font.size = Pt(10)
    r.font.name = 'Arial'
    return p


def _felsorolas(doc, szoveg):
    try:
        p = doc.add_paragraph(style='List Bullet')
    except Exception:
        p = doc.add_paragraph()
        p.add_run("• ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)

    if ':' in szoveg and len(szoveg.split(':', 1)[0]) < 40:
        elotag, maradek = szoveg.split(':', 1)
        r1 = p.add_run(elotag.strip() + ":")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = 'Arial'
        r2 = p.add_run(maradek)
        r2.font.size = Pt(10)
        r2.font.name = 'Arial'
    else:
        r = p.add_run(szoveg)
        r.font.size = Pt(10)
        r.font.name = 'Arial'
        def _cv_torzs(doc, cv_szoveg, akcent_hex):
    sorok = cv_szoveg.split('\n')
    i = 0
    n = len(sorok)
    while i < n:
        sor = sorok[i].strip()
        if not sor:
            i += 1
            continue

        if sor.startswith('## '):
            cim = sor[3:].strip()
            if 'PROFIL' in cim.upper():
                blokk = []
                i += 1
                while i < n and not sorok[i].strip().startswith('## '):
                    if sorok[i].strip() and not sorok[i].strip().startswith('---'):
                        blokk.append(sorok[i])
                    i += 1
                _profil_blokk(doc, blokk, akcent_hex)
                continue
            else:
                _szekcio_cim(doc, cim, akcent_hex)
        elif sor.startswith('# '):
            pass
        elif sor.startswith('---'):
            pass
        elif sor.startswith('**') and sor.endswith('**'):
            _munka_cim(doc, sor[2:-2])
        elif sor.startswith('*') and sor.endswith('*'):
            _munka_ceg(doc, sor[1:-1])
        elif sor.startswith('- ') or sor.startswith('✓ '):
            _felsorolas(doc, sor[2:].replace('**', ''))
        else:
            _bekezdes(doc, sor.replace('**', ''))
        i += 1


def _level_torzs(doc, level_szoveg):
    zaro_szavak = ("üdvözlettel", "tisztelettel", "köszönettel")
    for sor in level_szoveg.split('\n'):
        sor = sor.strip()
        if not sor:
            continue
        if sor.startswith('---'):
            continue

        tiszta = sor.replace('**', '')
        also = tiszta.lower()

        if (sor.startswith('**') and sor.endswith('**')) or also.startswith('tisztelt'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(tiszta)
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Arial'
            continue

        if any(also.startswith(z) for z in zaro_szavak):
            p = _bekezdes(doc, tiszta)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(2)
            continue

        if len(tiszta) < 40 and ' ' in tiszta and not tiszta.endswith('.'):
            p = _bekezdes(doc, tiszta)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            continue

        p = _bekezdes(doc, tiszta)
        p.paragraph_format.space_after = Pt(10)


def cv_docx_general(cv_szoveg: str, adatok: dict, szin_valasztas: str = "arany") -> bytes:
    szin = SZINEK.get(szin_valasztas, SZINEK["arany"])
    doc = _alap_dokumentum()
    _fejlec(doc, adatok, szin["fo"], szin["akcent"])
    _cv_torzs(doc, cv_szoveg, szin["akcent"])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def level_docx_general(level_szoveg: str, adatok: dict, szin_valasztas: str = "arany") -> bytes:
    szin = SZINEK.get(szin_valasztas, SZINEK["arany"])
    doc = _alap_dokumentum()
    _fejlec(doc, adatok, szin["fo"], szin["akcent"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _level_torzs(doc, level_szoveg)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
