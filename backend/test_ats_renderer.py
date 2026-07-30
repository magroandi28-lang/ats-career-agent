# -*- coding: utf-8 -*-
"""Az ATS Standard sablon szabályai (folyamat_terkep.md 8. és 11.10).

Ezek nem stílustesztek. Minden állítás mögött egy konkrét ok van, ami miatt
a robotszűrő elveszítheti a pályázót: fejlécbe zárt telefonszám, táblázatba
tördelt munkahely, beágyazott kép helyett kimaradt szöveg.
"""

from io import BytesIO

import pytest

from backend.ats_renderer import (
    DOCX_BETU,
    TORZS_PT,
    ats_docx,
    ats_pdf,
    szakaszokra_bontas,
)


MINTA_CV = """Kovács Anna
kovacs.anna@example.com
+36 30 123 4567
Budapest

Szakmai összefoglaló
Öt év bolti eladói tapasztalat, napi pénztárzárással.

SZAKMAI TAPASZTALAT
- Bolti eladó, ABC Kft. (2019-2024)
- Pénztáros, XY Bt. (2017-2019)

Tanulmányok:
Kereskedelmi szakközépiskola, érettségi

Nyelvek
angol, alapfok
"""


def _docx_dokumentum(tartalom: bytes):
    from docx import Document

    return Document(BytesIO(tartalom))


def test_szakaszokra_bontas_felismeri_a_valtozatokat():
    """A modell magyarul sokféleképpen nevezheti ugyanazt a szakaszt.

    Kettőspont, csupa nagybetű, ékezet -- ha ezeken elcsúsznánk, a szakasz
    tartalma a fejlécbe csúszna, és a CV tagolatlan tömbbé állna össze.
    """
    fejlec, szakaszok = szakaszokra_bontas(MINTA_CV)
    cimek = [cim for cim, _ in szakaszok]

    assert fejlec[0] == "Kovács Anna"
    assert "kovacs.anna@example.com" in fejlec
    assert cimek == [
        "Szakmai összefoglaló",
        "Szakmai tapasztalat",
        "Tanulmányok",
        "Nyelvek",
    ]


def test_hosszu_mondat_nem_lesz_szakaszcim():
    """Egy mondat, ami a „tapasztalat" szóval kezdődik, nem címsor."""
    szoveg = "Tapasztalatom van a napi pénztárzárásban és a készletkezelésben is."
    fejlec, szakaszok = szakaszokra_bontas(szoveg)

    assert szakaszok == []
    assert fejlec == [szoveg]


def test_ugyanaz_a_szakasz_ketszer_nem_kap_ket_cimsort():
    fejlec, szakaszok = szakaszokra_bontas(
        "Készségek\npénztárgép\n\nKompetenciák\nkészletkezelés"
    )

    assert [cim for cim, _ in szakaszok] == ["Készségek"]
    assert szakaszok[0][1] == ["pénztárgép", "készletkezelés"]


def test_ures_szakasz_nem_kap_cimsort():
    _, szakaszok = szakaszokra_bontas("Nyelvek\n\nKészségek\npénztárgép")

    assert [cim for cim, _ in szakaszok] == ["Készségek"]


def test_docx_nem_tartalmaz_tablazatot_es_kepet():
    """A spec pontról pontra tiltja. Táblázatból az ATS sorrendet téveszt."""
    dokumentum = _docx_dokumentum(ats_docx(MINTA_CV))

    assert dokumentum.tables == []
    # `inline_shapes` gyűjtemény, nem lista: a hosszát kell néznünk.
    assert len(dokumentum.inline_shapes) == 0
    # A beágyazott kép a dokumentum kapcsolatai közt is nyomot hagy.
    assert not [
        kapcsolat
        for kapcsolat in dokumentum.part.rels.values()
        if "image" in kapcsolat.reltype
    ]


def test_docx_nev_es_elerhetoseg_a_torzsben_van():
    """Nem fejlécben: azt sok ATS nem olvassa be, és a kapcsolat elveszne."""
    tartalom = ats_docx(MINTA_CV)
    dokumentum = _docx_dokumentum(tartalom)
    torzs = "\n".join(bekezdes.text for bekezdes in dokumentum.paragraphs)

    assert "Kovács Anna" in torzs
    assert "kovacs.anna@example.com" in torzs
    assert "+36 30 123 4567" in torzs

    fejlec_szoveg = "\n".join(
        bekezdes.text
        for szakasz in dokumentum.sections
        for bekezdes in szakasz.header.paragraphs
    )
    assert "kovacs.anna@example.com" not in fejlec_szoveg


def test_docx_betutipus_es_meret_a_spec_szerint():
    dokumentum = _docx_dokumentum(ats_docx(MINTA_CV))
    alap = dokumentum.styles["Normal"]

    assert alap.font.name == DOCX_BETU
    assert alap.font.size.pt == TORZS_PT
    assert 10.5 <= alap.font.size.pt <= 11


def test_docx_egyhasabos_marad():
    """Minden hasábolás ATS-kockázat: az olvasó összekeveri a sorrendet."""
    dokumentum = _docx_dokumentum(ats_docx(MINTA_CV))

    for szakasz in dokumentum.sections:
        oszlopok = szakasz._sectPr.xpath("./w:cols")
        for oszlop in oszlopok:
            darab = oszlop.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
            )
            assert darab in (None, "1")


def test_docx_minden_szakasz_tartalma_bekerul():
    dokumentum = _docx_dokumentum(ats_docx(MINTA_CV))
    torzs = "\n".join(bekezdes.text for bekezdes in dokumentum.paragraphs)

    assert "Bolti eladó, ABC Kft. (2019-2024)" in torzs
    assert "Kereskedelmi szakközépiskola, érettségi" in torzs
    assert "angol, alapfok" in torzs
    assert "SZAKMAI TAPASZTALAT" in torzs


def test_docx_nem_hasznal_szamozott_felsorolast():
    """A spec tiltja: a számozás az ATS-nek zaj, és sorrendet sugall."""
    dokumentum = _docx_dokumentum(ats_docx(MINTA_CV))

    for bekezdes in dokumentum.paragraphs:
        assert not bekezdes.style.name.lower().startswith("list number")


def test_pdf_elkeszul_es_valodi_pdf():
    tartalom = ats_pdf(MINTA_CV)

    assert tartalom.startswith(b"%PDF-")
    assert len(tartalom) > 500


def test_pdf_nem_bukik_el_a_kacsacsoron():
    """A reportlab mini-XML-ként olvassa a bekezdést.

    Egy „C&A" munkáltatónév enélkül értelmezési hibát dobna, és a CV EMIATT
    nem készülhetne el -- egy tipográfiai apróságon bukna a szolgáltatás.
    """
    tartalom = ats_pdf("Kovács Anna\n\nSzakmai tapasztalat\n- Eladó, C&A <Budapest>")

    assert tartalom.startswith(b"%PDF-")


def test_ures_szoveg_erthetoen_hibazik():
    for keszito in (ats_docx, ats_pdf):
        with pytest.raises(ValueError):
            keszito("   \n\n  ")


def test_cimsor_nelkuli_cv_is_elkeszul():
    """Ha a modell nem tagol, attól még legyen letölthető dokumentum."""
    dokumentum = _docx_dokumentum(ats_docx("Kovács Anna\nBolti eladó\nBudapest"))
    torzs = "\n".join(bekezdes.text for bekezdes in dokumentum.paragraphs)

    assert "Bolti eladó" in torzs


# ── A LETÖLTÉSI VÉGPONT ──────────────────────────────────────

from fastapi.testclient import TestClient  # noqa: E402

from backend.main import app  # noqa: E402
from backend.auth import jelenlegi_felhasznalo  # noqa: E402

kliens = TestClient(app)


class _Felhasznalo:
    id = "00000000-0000-0000-0000-000000000001"
    user_metadata = {}


def _workflow(context):
    return {"id": "workflow-1", "current_state": "CV_TERVEZET", "context": context}


def test_letoltes_a_tarolt_szovegbol_keszul(monkeypatch):
    """A drága lánc egyszer fut le; a letöltés ingyenes és kéréskor renderel."""
    from backend import main

    app.dependency_overrides[jelenlegi_felhasznalo] = lambda: _Felhasznalo()
    monkeypatch.setattr(main, "session_lekeres_vagy_letrehozas", lambda _: "session-1")
    monkeypatch.setattr(
        main,
        "workflow_lekeres_vagy_letrehozas",
        lambda *_: _workflow({"cv_uj_valtozat": MINTA_CV}),
    )
    try:
        valasz = kliens.post("/api/v1/cv/letoltes", json={"formatum": "pdf"})
    finally:
        app.dependency_overrides.clear()

    assert valasz.status_code == 200
    assert valasz.content.startswith(b"%PDF-")
    assert "oneletrajz.pdf" in valasz.headers["content-disposition"]


def test_letoltes_a_szerkesztett_valtozatot_hasznalja(monkeypatch):
    """Szerkesztés után a MÓDOSÍTOTT CV töltődik le, nem a tárolt eredeti."""
    from backend import main

    app.dependency_overrides[jelenlegi_felhasznalo] = lambda: _Felhasznalo()
    monkeypatch.setattr(main, "session_lekeres_vagy_letrehozas", lambda _: "session-1")
    monkeypatch.setattr(
        main,
        "workflow_lekeres_vagy_letrehozas",
        lambda *_: _workflow({"cv_uj_valtozat": MINTA_CV}),
    )
    try:
        valasz = kliens.post(
            "/api/v1/cv/letoltes",
            json={
                "formatum": "docx",
                "cv_szoveg": "Szabó Béla\n\nKészségek\nraktári adminisztráció",
            },
        )
    finally:
        app.dependency_overrides.clear()

    assert valasz.status_code == 200
    dokumentum = _docx_dokumentum(valasz.content)
    torzs = "\n".join(bekezdes.text for bekezdes in dokumentum.paragraphs)
    assert "Szabó Béla" in torzs
    assert "Kovács Anna" not in torzs


def test_letoltes_elkeszult_cv_nelkul_erthetoen_elutasit(monkeypatch):
    """Ne ures fajlt adjunk: mondjuk meg, hogy elobb le kell futnia a lancnak."""
    from backend import main

    app.dependency_overrides[jelenlegi_felhasznalo] = lambda: _Felhasznalo()
    monkeypatch.setattr(main, "session_lekeres_vagy_letrehozas", lambda _: "session-1")
    monkeypatch.setattr(
        main, "workflow_lekeres_vagy_letrehozas", lambda *_: _workflow({})
    )
    try:
        valasz = kliens.post("/api/v1/cv/letoltes", json={"formatum": "docx"})
    finally:
        app.dependency_overrides.clear()

    assert valasz.status_code == 409


def test_letoltes_ismeretlen_formatumot_nem_fogad_el(monkeypatch):
    from backend import main

    app.dependency_overrides[jelenlegi_felhasznalo] = lambda: _Felhasznalo()
    monkeypatch.setattr(main, "session_lekeres_vagy_letrehozas", lambda _: "session-1")
    monkeypatch.setattr(
        main,
        "workflow_lekeres_vagy_letrehozas",
        lambda *_: _workflow({"cv_uj_valtozat": MINTA_CV}),
    )
    try:
        valasz = kliens.post("/api/v1/cv/letoltes", json={"formatum": "rtf"})
    finally:
        app.dependency_overrides.clear()

    assert valasz.status_code == 422


def test_letoltes_bejelentkezes_nelkul_nem_elerheto():
    valasz = kliens.post("/api/v1/cv/letoltes", json={"formatum": "pdf"})

    assert valasz.status_code == 401
