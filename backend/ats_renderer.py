# -*- coding: utf-8 -*-
"""ATS Standard sablon — a CV megjelenését KÓD állítja elő, nem a modell.

A `folyamat_terkep.md` 8. fejezete egyetlen sablont ír elő az első kiadásra.
A lényeg nem a szépség, hanem hogy a robotszűrő ki tudja olvasni:

- egy hasáb, táblázat és szövegdoboz nélkül;
- Arial vagy Calibri, 10,5-11 pontos törzsszöveg;
- a név és az elérhetőség a dokumentum TÖRZSÉBEN, nem fejlécben vagy
  láblécben (a fejlécet sok ATS egyszerűen nem olvassa be);
- nincs fotó, ikon, diagram, táblázat;
- egységes címsorok.

MIÉRT NEM A MEGLÉVŐ SABLONOKAT HASZNÁLJUK. A `utils/docx_sablonok.py` és a
`utils/pdf_sablonok.py` a Streamlit-korszakból maradt, és pontosan azt teszi,
amit ez a fejezet tilt: kör alakú fotót, telefon- és e-mail-ikont, akcentszínt,
szegélyt és betűritkítást. Azokat nem átszabni kellett volna, hanem
mellétenni ezt -- ATS-re a díszítés nem semleges, hanem kockázat.

A modell szabad szöveges CV-t ad vissza (`CvFactCheck.checked_cv`). Ezt itt
determinisztikusan bontjuk szakaszokra: ugyanaz a bemenet mindig ugyanazt a
kimenetet adja, tehát a CV kinézete nem függ a modell pillanatnyi kedvétől.
"""

from __future__ import annotations

import re
import unicodedata
from io import BytesIO


# A spec 8. fejezetének szakaszai, ebben a sorrendben. A kulcs a kanonikus
# cím, az érték a felismert változatok -- a modell magyarul sokféleképpen
# nevezheti ugyanazt, és nem akarunk emiatt szakaszt elveszíteni.
SZAKASZOK: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "Szakmai összefoglaló",
        (
            "szakmai osszefoglalo", "osszefoglalo", "szakmai profil", "profil",
            "bemutatkozas", "rolam", "szakmai bemutatkozas",
        ),
    ),
    (
        "Szakmai tapasztalat",
        (
            "szakmai tapasztalat", "tapasztalat", "munkatapasztalat",
            "munkahelyek", "szakmai eletut", "korabbi munkahelyek",
        ),
    ),
    (
        "Tanulmányok",
        (
            "tanulmanyok", "vegzettseg", "vegzettsegek", "kepzettseg",
            "iskolai vegzettseg", "oktatas", "tanulmanyi hatter",
        ),
    ),
    (
        "Készségek",
        (
            "keszsegek", "kompetenciak", "szakmai keszsegek", "tudas",
            "ismeretek", "keszseg", "szakmai kompetenciak",
        ),
    ),
    (
        "Nyelvek",
        ("nyelvek", "nyelvtudas", "nyelvismeret", "idegen nyelvek"),
    ),
    (
        "Projektek",
        ("projektek", "projekt", "referenciak", "portfolio", "munkak"),
    ),
)

# Törzsszöveg mérete pontban. A spec 10,5-11 pontot enged; a 11 a felső
# határ, mert a kisebb betű nyomtatásban és képernyőolvasóval is nehezebb.
TORZS_PT = 11
CIM_PT = 12
NEV_PT = 16

# A DOCX-ben Calibri, a PDF-ben Helvetica. Mindkettő megfelel a spec
# „Arial vagy Calibri" kikötésének: a Helvetica az Arial szabványos PDF-beli
# megfelelője, és a PDF alapkészletének tagja -- nem kell beágyazni, tehát
# nem tud hiányzó betűtípus miatt szétesni egy másik gépen.
DOCX_BETU = "Calibri"
PDF_BETU = "Helvetica"
PDF_BETU_FELKOVER = "Helvetica-Bold"

# Felsorolásjelek, amiket a modell használhat. Egységesítjük, mert a vegyes
# jelölés a robotszűrőnek zaj, az olvasónak meg rendetlenség.
_FELSOROLAS_ELE = re.compile(r"^\s*(?:[-*•–—·o]|\d+[.)])\s+")


def _ekezet_nelkul(szoveg: str) -> str:
    """Kisbetűs, ékezet nélküli alak — csak összehasonlításhoz."""

    bontott = unicodedata.normalize("NFKD", szoveg or "")
    return "".join(jel for jel in bontott if not unicodedata.combining(jel)).casefold()


def _szakaszcim(sor: str) -> str | None:
    """A sor kanonikus szakaszcíme, ha az. Különben None.

    A modell írhat kettőspontot, csupa nagybetűt vagy `##` jelölést -- ezek
    mind ugyanazt a szakaszt jelentik. A hosszú sorokat viszont NEM tekintjük
    címnek: egy mondat, ami történetesen a „tapasztalat" szóval kezdődik, nem
    szakaszcím, hanem tartalom.
    """

    tisztitott = (sor or "").strip().strip("#").strip()
    tisztitott = tisztitott.rstrip(":").strip()
    if not tisztitott or len(tisztitott) > 40:
        return None

    kulcs = _ekezet_nelkul(tisztitott)
    for cim, valtozatok in SZAKASZOK:
        if kulcs in valtozatok:
            return cim
    return None


def szakaszokra_bontas(cv_szoveg: str) -> tuple[list[str], list[tuple[str, list[str]]]]:
    """A szabad szöveges CV-t fejlécre és szakaszokra bontja.

    Visszaad: (fejléc sorai, [(szakaszcím, sorok)]).

    A fejléc az első felismert szakaszcím ELŐTTI rész: itt áll a név és az
    elérhetőség. Ez szándékosan a dokumentum törzsébe kerül -- a Word-fejlécet
    sok ATS nem olvassa be, és pont a kapcsolati adat veszne el.

    Ha a modell egyáltalán nem használ felismerhető címet, minden a fejlécbe
    kerül. Ez nem hiba: attól még olvasható CV, csak tagolatlan.
    """

    fejlec: list[str] = []
    szakaszok: list[tuple[str, list[str]]] = []
    aktualis: list[str] | None = None

    for nyers in (cv_szoveg or "").splitlines():
        sor = nyers.rstrip()
        cim = _szakaszcim(sor)
        if cim is not None:
            # Ugyanaz a szakasz kétszer: ne keletkezzen két azonos címsor.
            meglevo = next((tetel for tetel in szakaszok if tetel[0] == cim), None)
            if meglevo is not None:
                aktualis = meglevo[1]
            else:
                aktualis = []
                szakaszok.append((cim, aktualis))
            continue
        if not sor.strip():
            continue
        (aktualis if aktualis is not None else fejlec).append(sor.strip())

    # Üres szakasz nem kap címsort: egy „Nyelvek" felirat alatta semmivel
    # rosszabb, mint ha ott sem lenne.
    return fejlec, [(cim, sorok) for cim, sorok in szakaszok if sorok]


def _felsorolas_e(sor: str) -> tuple[bool, str]:
    """Felsorolás-e a sor, és mi a szövege a jelölés nélkül."""

    ujra = _FELSOROLAS_ELE.sub("", sor)
    return (ujra != sor, ujra.strip() or sor.strip())


def ats_docx(cv_szoveg: str) -> bytes:
    """ATS Standard DOCX. Egy hasáb, táblázat és kép nélkül.

    A DOCX az elsődleges szerkeszthető fájl (spec 8.), ezért ez a mérvadó
    kimenet; a PDF ugyanezt rögzíti.
    """

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    fejlec, szakaszok = szakaszokra_bontas(cv_szoveg)
    if not fejlec and not szakaszok:
        raise ValueError("Üres CV-szövegből nem készíthető dokumentum.")

    dokumentum = Document()

    # Egy hasáb, mérsékelt margó. A `Document()` alapból egyhasábos -- ezt
    # szándékosan nem bántjuk, mert minden hasábolás ATS-kockázat.
    for szakasz in dokumentum.sections:
        szakasz.top_margin = Cm(1.8)
        szakasz.bottom_margin = Cm(1.8)
        szakasz.left_margin = Cm(2.0)
        szakasz.right_margin = Cm(2.0)

    alap = dokumentum.styles["Normal"]
    alap.font.name = DOCX_BETU
    alap.font.size = Pt(TORZS_PT)

    # A NÉV ÉS AZ ELÉRHETŐSÉG A TÖRZSBEN.
    #
    # Nem `section.header`-be: a Word-fejlécet sok ATS nem olvassa be, és
    # akkor pont a kapcsolati adat tűnik el a pályázatból.
    if fejlec:
        nev = dokumentum.add_paragraph()
        nev.alignment = WD_ALIGN_PARAGRAPH.LEFT
        futas = nev.add_run(fejlec[0])
        futas.bold = True
        futas.font.size = Pt(NEV_PT)
        for sor in fejlec[1:]:
            bekezdes = dokumentum.add_paragraph(sor)
            bekezdes.paragraph_format.space_after = Pt(0)

    for cim, sorok in szakaszok:
        cimsor = dokumentum.add_paragraph()
        cimsor.paragraph_format.space_before = Pt(12)
        cimsor.paragraph_format.space_after = Pt(4)
        futas = cimsor.add_run(cim.upper())
        futas.bold = True
        futas.font.size = Pt(CIM_PT)

        for sor in sorok:
            felsorolas, szoveg = _felsorolas_e(sor)
            # `List Bullet` helyett kézi jel: a listastílus a Word verziójától
            # függően számozott listává alakulhat, a spec pedig tiltja a
            # számozott felsorolást. Egy sima kötőjel minden olvasóban ugyanaz.
            bekezdes = dokumentum.add_paragraph(f"– {szoveg}" if felsorolas else szoveg)
            bekezdes.paragraph_format.space_after = Pt(2)
            if felsorolas:
                bekezdes.paragraph_format.left_indent = Cm(0.5)

    puffer = BytesIO()
    dokumentum.save(puffer)
    return puffer.getvalue()


def ats_pdf(cv_szoveg: str) -> bytes:
    """ATS Standard PDF — ugyanaz az elrendezés, rögzített megjelenéssel."""

    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    fejlec, szakaszok = szakaszokra_bontas(cv_szoveg)
    if not fejlec and not szakaszok:
        raise ValueError("Üres CV-szövegből nem készíthető dokumentum.")

    nev_stilus = ParagraphStyle(
        "AtsNev", fontName=PDF_BETU_FELKOVER, fontSize=NEV_PT,
        leading=NEV_PT + 4, alignment=TA_LEFT, spaceAfter=4,
    )
    torzs_stilus = ParagraphStyle(
        "AtsTorzs", fontName=PDF_BETU, fontSize=TORZS_PT,
        leading=TORZS_PT + 4, alignment=TA_LEFT, spaceAfter=2,
    )
    cim_stilus = ParagraphStyle(
        "AtsCim", fontName=PDF_BETU_FELKOVER, fontSize=CIM_PT,
        leading=CIM_PT + 3, alignment=TA_LEFT, spaceBefore=10, spaceAfter=4,
    )
    felsorolas_stilus = ParagraphStyle(
        "AtsFelsorolas", parent=torzs_stilus, leftIndent=0.5 * cm,
    )

    elemek = []
    if fejlec:
        elemek.append(Paragraph(_pdf_biztos(fejlec[0]), nev_stilus))
        for sor in fejlec[1:]:
            elemek.append(Paragraph(_pdf_biztos(sor), torzs_stilus))

    for cim, sorok in szakaszok:
        elemek.append(Paragraph(_pdf_biztos(cim.upper()), cim_stilus))
        for sor in sorok:
            felsorolas, szoveg = _felsorolas_e(sor)
            elemek.append(
                Paragraph(
                    _pdf_biztos(f"– {szoveg}" if felsorolas else szoveg),
                    felsorolas_stilus if felsorolas else torzs_stilus,
                )
            )
    elemek.append(Spacer(1, 2))

    puffer = BytesIO()
    SimpleDocTemplate(
        puffer, pagesize=A4,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm,
        leftMargin=2.0 * cm, rightMargin=2.0 * cm,
        title="Önéletrajz", author="",
    ).build(elemek)
    return puffer.getvalue()


def _pdf_biztos(szoveg: str) -> str:
    """A reportlab a bekezdésszöveget mini-XML-ként értelmezi.

    Egy `<` vagy `&` a CV-ben (például „C&A" munkáltatónév) enélkül
    értelmezési hibát dobna, és a CV-nek EMIATT nem lehetne elkészülnie --
    egy tipográfiai apróság miatt bukna a szolgáltatás.
    """

    return (
        (szoveg or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
