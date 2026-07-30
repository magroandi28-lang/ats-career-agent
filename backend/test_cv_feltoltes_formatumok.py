# -*- coding: utf-8 -*-
"""A CV-feltöltés elfogadott formátumai (folyamat_terkep.md 2.3 és 11.2).

A felület CSAK olyat kínálhat fel, amit a backend tényleg fel is dolgoz.
Ezért a formátumellenőrzést és a szövegkinyerést együtt mérjük: külön-külön
mindkettő zöld lehetne úgy, hogy a felhasználónál mégis elakad a feltöltés.
"""

import asyncio
from io import BytesIO

import pytest
from fastapi import UploadFile
from starlette.datastructures import Headers

from backend.cv_import_service import extract_cv_text, extract_docx_text
from backend.security import read_validated_cv_file


def _feltoltes(tartalom: bytes, filename: str, content_type: str) -> UploadFile:
    return UploadFile(
        BytesIO(tartalom),
        filename=filename,
        headers=Headers({"content-type": content_type}),
    )


def _docx_bytes(bekezdesek=(), tablazat_sorok=()) -> bytes:
    """Valódi .docx fájl, nem utánzat -- a python-docx sajátjától tanulunk."""
    from docx import Document

    dokumentum = Document()
    for szoveg in bekezdesek:
        dokumentum.add_paragraph(szoveg)
    if tablazat_sorok:
        tabla = dokumentum.add_table(rows=len(tablazat_sorok), cols=2)
        for sor_index, (bal, jobb) in enumerate(tablazat_sorok):
            tabla.rows[sor_index].cells[0].text = bal
            tabla.rows[sor_index].cells[1].text = jobb
    puffer = BytesIO()
    dokumentum.save(puffer)
    return puffer.getvalue()


PNG_FEJLEC = b"\x89PNG\r\n\x1a\n" + b"teszt"
JPG_FEJLEC = b"\xff\xd8\xff" + b"teszt"


@pytest.mark.parametrize(
    "eset",
    ["pdf", "docx", "jpg", "jpeg", "png"],
)
def test_mind_a_negy_formatum_elfogadva(eset):
    """A spec négy formátumot ígér a felhasználónak. Mind a négy menjen át.

    A paraméter szándékosan rövid név, nem maga a fájl: a bináris DOCX
    beleírná magát a teszt azonosítójába, és egyetlen hibaüzenet
    olvashatatlanná válna tőle.
    """
    esetek = {
        "pdf": (b"%PDF-1.7\nteszt", "cv.pdf", "application/pdf", "pdf"),
        "docx": (
            _docx_bytes(("Teszt",)),
            "cv.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        ),
        "jpg": (JPG_FEJLEC, "cv.jpg", "image/jpeg", "kep"),
        "jpeg": (JPG_FEJLEC, "cv.jpeg", "image/jpeg", "kep"),
        "png": (PNG_FEJLEC, "cv.png", "image/png", "kep"),
    }
    tartalom, filename, mime, vart_formatum = esetek[eset]

    olvasott, formatum = asyncio.run(
        read_validated_cv_file(_feltoltes(tartalom, filename, mime))
    )
    assert olvasott == tartalom
    assert formatum == vart_formatum


def test_altalanos_mime_nem_bukik_el():
    """Több rendszer application/octet-stream-et küld minden csatolmányra.

    A kiterjesztés és a magic byte ellenőrzés ilyenkor is érvényben marad,
    tehát ez nem lyuk -- viszont enélkül a felhasználó feltöltése érthetetlen
    okból akadna el, pedig a fájlja rendben van.
    """
    tartalom = _docx_bytes(("Teszt",))
    _, formatum = asyncio.run(
        read_validated_cv_file(
            _feltoltes(tartalom, "cv.docx", "application/octet-stream")
        )
    )
    assert formatum == "docx"


def test_hamis_kiterjesztes_a_magic_byte_on_bukik():
    """A kiterjesztést a támadó állítja, a magic byte-ot nem."""
    with pytest.raises(Exception) as hiba:
        asyncio.run(
            read_validated_cv_file(
                _feltoltes(
                    b"MZ\x90\x00futtathato",
                    "cv.docx",
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document",
                )
            )
        )
    assert getattr(hiba.value, "status_code", None) == 400


def test_nem_tamogatott_formatum_elutasitva():
    with pytest.raises(Exception) as hiba:
        asyncio.run(
            read_validated_cv_file(_feltoltes(b"barmi", "cv.doc", "application/msword"))
        )
    assert getattr(hiba.value, "status_code", None) == 400


def test_docx_tablazatbol_is_olvas():
    """A magyar CV-sablonok többsége táblázatba tördeli a munkahelyeket.

    Ha csak a bekezdéseket néznénk, pont a szakmai előzmény veszne el, és a
    felhasználó azt látná: „nincs kinyerhető szöveg".
    """
    tartalom = _docx_bytes(
        bekezdesek=("Kovács Anna", "Bolti eladó"),
        tablazat_sorok=(("2019-2024", "Bolti eladó, ABC Kft."),),
    )
    szoveg = extract_docx_text(tartalom)

    assert "Kovács Anna" in szoveg
    assert "ABC Kft." in szoveg
    assert "2019-2024" in szoveg


def test_ures_docx_erthetoen_hibazik():
    with pytest.raises(ValueError) as hiba:
        extract_docx_text(_docx_bytes())
    assert "kinyerhető szöveget" in str(hiba.value)


def test_doc_fajl_erthetoen_iranyit_tovabb():
    """A .doc nem ZIP, tehát a python-docx nem tudja megnyitni.

    A hibaüzenet mondja meg, mit tegyen -- ne csak azt, hogy nem sikerült.
    """
    with pytest.raises(ValueError) as hiba:
        extract_docx_text(b"\xd0\xcf\x11\xe0regi word")
    assert ".docx" in str(hiba.value)


def test_kep_kinyerese_az_atirast_hasznalja(monkeypatch):
    """Képnél nincs mit determinisztikusan olvasni -- ez az egyetlen modelles út."""
    import utils.flow_agy as flow_agy

    monkeypatch.setattr(flow_agy, "kep_atiras", lambda *_: "Kovács Anna\nBolti eladó")
    assert "Kovács Anna" in extract_cv_text(PNG_FEJLEC, "kep", "image/png")


def test_olvashatatlan_kep_erthetoen_hibazik(monkeypatch):
    import utils.flow_agy as flow_agy

    monkeypatch.setattr(flow_agy, "kep_atiras", lambda *_: "")
    with pytest.raises(ValueError) as hiba:
        extract_cv_text(PNG_FEJLEC, "kep", "image/png")
    assert "PDF" in str(hiba.value)
